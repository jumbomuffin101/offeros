from app.core.errors import AppError, ValidationError
from app.services.resume_extraction.base import ExtractionResult, MAX_EXTRACTED_TEXT_CHARS
from app.services.resume_extraction.utils import normalize_resume_text


class DocxResumeExtractor:
    def extract(self, content: bytes) -> ExtractionResult:
        try:
            from docx import Document
        except ImportError as exc:
            raise AppError("extraction_not_configured", "DOCX extraction is not configured on the backend.", 503) from exc

        try:
            from io import BytesIO
            document = Document(BytesIO(content))
        except Exception as exc:
            raise ValidationError("OfferOS could not read this DOCX file.") from exc

        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        normalized = normalize_resume_text("\n".join(lines))
        if len(normalized) < 80:
            raise ValidationError("This DOCX file does not contain enough readable resume text.")
        warnings: list[str] = []
        if len(normalized) > MAX_EXTRACTED_TEXT_CHARS:
            warnings.append("Extracted text was truncated to the maximum supported length.")
            normalized = normalized[:MAX_EXTRACTED_TEXT_CHARS]
        return ExtractionResult(text=normalized, page_count=None, warnings=warnings)
